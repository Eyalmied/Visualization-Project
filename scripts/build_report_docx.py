"""
Builds report/Report.docx from report/stats.json and report/figures/*.png.
Re-run after `Rscript scripts/generate_report_assets.R` if the data,
charts, or numbers change.

Usage: python scripts/build_report_docx.py
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "report" / "figures"
STATS = json.loads((ROOT / "report" / "stats.json").read_text())

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x66, 0x66, 0x66)


def fmt(n):
    return f"{n:,}"


def pct(row_list, key, value):
    for r in row_list:
        if r[key] == value:
            return r
    return None


doc = Document()

# ---- base style ------------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for i, size in zip(range(1, 4), (20, 16, 13)):
    hstyle = doc.styles[f"Heading {i}"]
    hstyle.font.color.rgb = ACCENT
    hstyle.font.size = Pt(size)


def placeholder(text):
    p = doc.add_paragraph()
    run = p.add_run(f"[PLACEHOLDER — {text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    return p


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_figure(filename, cap, width=6.0):
    doc.add_picture(str(FIG / filename), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(cap)


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Indian Flights Dashboard")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("An Interactive Visual Analytics Dashboard for Flight Price Exploration")
run.font.size = Pt(14)
run.font.color.rgb = GREY

doc.add_paragraph()
course = doc.add_paragraph()
course.alignment = WD_ALIGN_PARAGRAPH.CENTER
course.add_run("Visualization of Information — Course Project Report").bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "Dataset: airlines_flights_data.csv (Kaggle “Flight Price Prediction”, Indian domestic flights)\n"
    "Repository: https://github.com/Eyalmied/Visualization-Project\n"
)

doc.add_paragraph()
placeholder("Team member names and roles (see PLAN.md → Team Task Breakdown for the 4 assigned roles: "
            "Data Lead, Traveler-View Lead, Airline-Analytics Lead, Integration & Report Lead)")
placeholder("Submission date / presentation slot used (2026-07-15 18:00–20:00 or 2026-07-16 20:00–21:00)")

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS (manual, since python-docx can't build a live TOC field
# reliably without the user updating fields in Word)
# ============================================================================
doc.add_heading("Contents", level=1)
toc_items = [
    "1. Problem Definition",
    "2. Data Understanding",
    "3. Data Processing",
    "4. Visualization & Interaction Design",
    "5. AI Usage Documentation",
    "6. Evaluation & Reflection",
    "7. Deliverables Checklist",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph()
note = doc.add_paragraph()
note.add_run("Note: ").bold = True
note.add_run("In Word, right-click this list and choose “Update Field”, or replace it with "
             "References → Table of Contents once section pages are finalized.")

doc.add_page_break()

# ============================================================================
# 1. PROBLEM DEFINITION
# ============================================================================
doc.add_heading("1. Problem Definition", level=1)

doc.add_heading("1.1 The analytics problem", level=2)
doc.add_paragraph(
    "Choosing a domestic flight in India means navigating six airlines, six major cities, three stop-count "
    "tiers, six departure/arrival time bands, and a booking window that runs from 49 days out to the day of "
    "travel — with prices that can differ by two orders of magnitude (₹{} to ₹{}) depending on the "
    "combination chosen. Neither a single summary statistic nor a static chart can answer a concrete question "
    "like “when should I book my Delhi→Mumbai flight, and on which airline?” because the answer "
    "depends on which of these dimensions the traveler actually cares about, and those dimensions interact "
    "(see §2.3). The analytics problem this project addresses is: ".format(
        fmt(STATS["price_min"]), fmt(STATS["price_max"]))
    + "how can a decision-maker explore a 300,000-row, multi-dimensional fare dataset "
    "interactively enough to find a trustworthy, personally-relevant answer, rather than relying on a single "
    "misleading average?"
)

doc.add_heading("1.2 Target users and decision-support goals", level=2)
add_table(
    ["Persona", "Decision-support goal", "Dashboard tab"],
    [
        ("Traveler", "Find the cheapest / best-value combination of route, airline, class, stop count, and "
                      "booking lead time for an upcoming trip.", "Traveler View"),
        ("Airline / revenue analyst", "Benchmark market share, average pricing, and route performance across "
                                       "the six competing carriers to inform pricing or route strategy.", "Airline Analytics"),
    ],
)

doc.add_heading("1.3 How the interactive app solves the problem", level=2)
doc.add_paragraph(
    "The dashboard deliberately does not use a predictive model (see §6.2 for the rationale). Instead, "
    "decision support is delivered entirely through filtering and comparison: every chart on both tabs is "
    "driven by one shared, reactive, filtered dataset per tab, so a user narrows the data to exactly their "
    "situation (e.g. “Economy, Delhi→Mumbai, booking within 10 days”) and every chart, and every "
    "KPI tile, updates to reflect that specific slice — rather than forcing them to mentally re-derive an "
    "answer from an aggregate chart that mixes together situations that don't apply to them. This directly "
    "implements Shneiderman's Visual Information Seeking Mantra (“overview first, zoom/filter, "
    "details-on-demand”, Lecture 2) and Keim's extended version (Lecture 5): the KPI tiles give the "
    "overview, the sidebar filters do the zooming, and the charts plus hover tooltips give the detail."
)

doc.add_page_break()

# ============================================================================
# 2. DATA UNDERSTANDING
# ============================================================================
doc.add_heading("2. Data Understanding", level=1)

doc.add_heading("2.1 Source and structure", level=2)
doc.add_paragraph(
    "The dataset is airlines_flights_data.csv, the Kaggle “Flight Price Prediction” dataset covering "
    f"Indian domestic flights. It contains {fmt(STATS['n_rows'])} rows and 12 raw columns: index, airline, "
    "flight, source_city, departure_time, stops, arrival_time, destination_city, class, duration, days_left, "
    "and price. There is no real calendar date (only days_left, an integer 1–49 representing booking lead "
    "time) and no latitude/longitude — only city names for the six cities served "
    f"({STATS['n_cities']} cities, {STATS['n_routes']} unique ordered route pairs, {STATS['n_airlines']} airlines, "
    f"{fmt(1561)} distinct flight codes). A full NA and duplicate-row check found "
    f"{STATS['n_na']} missing values and, once the meaningless row-number index column is excluded, "
    f"{fmt(STATS['n_rows'])} of {fmt(300153)} rows are unique (2 exact duplicate listings were removed — "
    "see §3.1)."
)

doc.add_heading("2.2 Overall price and duration profile", level=2)
add_table(
    ["Metric", "Value"],
    [
        ("Price range", f"₹{fmt(STATS['price_min'])} – ₹{fmt(STATS['price_max'])}"),
        ("Price mean / median", f"₹{fmt(STATS['price_mean'])} / ₹{fmt(STATS['price_median'])}"),
        ("Price standard deviation", f"₹{fmt(STATS['price_sd'])}"),
        ("Duration range", f"{STATS['duration_min']} – {STATS['duration_max']} hours"),
        ("Duration mean", f"{STATS['duration_mean']} hours"),
        ("Correlation (duration, price)", f"{STATS['duration_price_cor']} (weak positive)"),
    ],
)
doc.add_paragraph(
    "The mean price (₹{}) is nearly three times the median (₹{}) — a classic right-skew signature. "
    "This is the direct, data-driven justification for the log-scale toggle built into chart #1 (§4.2) and "
    "for citing Lecture 4's log-transform guidance rather than treating it as a generic best practice.".format(
        fmt(STATS["price_mean"]), fmt(STATS["price_median"]))
)

doc.add_heading("2.3 A genuine Anscombe's-Quartet-style finding: aggregate price by airline is misleading", level=2)
doc.add_paragraph(
    "A first pass at “average price by airline” produces a striking, almost bimodal split:"
)
add_table(
    ["Airline", "Mean price", "Flights", "Market share"],
    [
        (r["airline"], f"₹{fmt(round(r['mean_price']))}", fmt(r["n"]),
         f"{pct(STATS['market_share'], 'airline', r['airline'])['share_pct']}%")
        for r in STATS["price_by_airline"]
    ],
)
doc.add_paragraph(
    "Vistara and Air_India appear roughly 4–7× more expensive than the other four carriers. The "
    "instinctive reading — “Vistara and Air_India are the premium/expensive airlines” — is "
    "exactly the kind of conclusion Anscombe's Quartet (Lecture 1) warns against: a single aggregate statistic "
    "can look meaningful while hiding the actual structure of the data. Checking the class mix per airline "
    "reveals the real driver: AirAsia, GO_FIRST, Indigo, and SpiceJet operate Economy-only in this dataset, "
    "while Vistara and Air_India are the only two carriers that also sell Business class (47.4% and 40.7% of "
    "their flights respectively). Once price is conditioned on class, the picture changes substantially:"
)
add_table(
    ["Airline", "Class", "Mean price", "Flights"],
    [
        ("Air_India", "Economy", "₹7,314", "47,994"),
        ("Air_India", "Business", "₹47,131", "32,898"),
        ("Vistara", "Economy", "₹7,807", "67,270"),
        ("Vistara", "Business", "₹55,477", "60,589"),
        ("AirAsia", "Economy", "₹4,091", "16,098"),
        ("GO_FIRST", "Economy", "₹5,652", "23,173"),
        ("Indigo", "Economy", "₹5,324", "43,120"),
        ("SpiceJet", "Economy", "₹6,179", "9,011"),
    ],
)
doc.add_paragraph(
    "Within Economy, Vistara and Air_India are only modestly pricier than the budget carriers (₹7,300–"
    "₹7,800 vs ₹4,000–₹6,200) — the ~4–7× gap in the unconditioned averages is "
    "overwhelmingly a class-mix artifact, not a fair “which airline is expensive” comparison. The same "
    "pattern recurs with the stops column: the raw mean price by stop count (zero: ₹9,376; one: ₹22,901; "
    "two_or_more: ₹14,113) looks like one-stop flights are the most expensive, which is counter-intuitive. "
    "Checking Business-class share by stop count explains it: 33.6% of one-stop flights are Business class, "
    "versus 22.5% of nonstop and only 8.2% of two-or-more-stop flights — again, class mix, not stop count "
    "itself, is the dominant driver of the aggregate. This is precisely why the dashboard's class checkbox filter "
    "is not optional decoration: without it, a user comparing airlines or stop counts would draw the wrong "
    "conclusion from a chart that looks perfectly reasonable."
)

doc.add_heading("2.4 Booking-lead-time trend", level=2)
doc.add_paragraph(
    "Averaging across all flights, booking within the last 3 days before departure costs ₹{} on average, "
    "versus ₹{} when booking 30+ days out — a {}% premium for last-minute booking. This trend, and how "
    "it varies by class and airline, is the core content of chart #4 (§4.2).".format(
        fmt(STATS["booking_early_mean"]), fmt(STATS["booking_late_mean"]), STATS["booking_pct_diff"])
)

doc.add_heading("2.5 Route-level patterns", level=2)
add_table(
    ["Most expensive routes (avg.)", "Cheapest routes (avg.)", "Busiest routes (by flight count)"],
    [
        (f"{a['route']}: ₹{fmt(round(a['mean_price']))}",
         f"{b['route']}: ₹{fmt(round(b['mean_price']))}",
         f"{c['route']}: {fmt(c['n'])} flights")
        for a, b, c in zip(STATS["top_expensive_routes"], STATS["top_cheap_routes"], STATS["top_busy_routes"])
    ],
)
doc.add_paragraph(
    "The busiest routes (Delhi↔Mumbai, Delhi→Bangalore) are not the cheapest or most expensive on "
    "average — route price and route volume are largely independent, which is exactly what chart #5's "
    "route heatmap and chart #6's route map are designed to let a user discover for themselves rather than "
    "read off a single ranked list."
)

doc.add_page_break()

# ============================================================================
# 3. DATA PROCESSING
# ============================================================================
doc.add_heading("3. Data Processing", level=1)

doc.add_heading("3.1 Cleaning", level=2)
doc.add_paragraph(
    "The raw index column is dropped (it is a row-number artifact, not analytical data). NA and duplicate "
    "checks are run and logged on every load: 0 missing values were found; 2 exact duplicate rows (identical "
    "across every column except the meaningless index) were removed, taking the working dataset from 300,153 "
    f"to {fmt(STATS['n_rows'])} rows. Both checks are logged (not silently assumed) so the Data Understanding "
    "claim of “clean data” in §2.1 is verifiable, not asserted."
)

doc.add_heading("3.2 Factor ordering", level=2)
doc.add_paragraph(
    "departure_time, arrival_time, and stops arrive as plain strings that would sort alphabetically on a chart "
    "axis (e.g. “Afternoon, Early_Morning, Evening, Late_Night, Morning, Night”), misrepresenting the "
    "data's real order. Each is converted to an ordered factor with the correct real-world sequence "
    "(Early_Morning → Morning → Afternoon → Evening → Night → Late_Night for time-of-day; "
    "zero → one → two_or_more for stops). This is a direct application of Lecture 2's “order” "
    "property of visual variables: an axis that doesn't reflect the data's natural order is a graphical-integrity "
    "problem, not a cosmetic one — chart #8 (departure-time × airline grid) would be actively misleading "
    "without this fix."
)

doc.add_heading("3.3 Feature engineering", level=2)
add_table(
    ["Feature", "Definition", "Why it was needed"],
    [
        ("price_per_hour", "price / duration",
         "Neither column alone lets a traveler compare “value” across a short vs. a long route, or an "
         "airline analyst compare pricing efficiency across carriers; this normalizes for flight length."),
        ("route", "paste(source_city, \"→\", destination_city)",
         "Collapses two categorical columns into one for route-level charts (heatmap, KPI tiles) without which "
         "grouping by “route” would require a two-column group_by everywhere it's used."),
        ("days_left_bucket", "cut into Last-minute (1–7) / 2–4 weeks (8–28) / 1+ month (29–49)",
         "49 raw integer values is too granular for a clean group comparison; kept as a derived column while the "
         "raw days_left is retained for the continuous slider filter and the booking-curve chart."),
        ("src_lat / src_lon / dst_lat / dst_lon", "joined from a hand-built 6-city lookup table",
         "The raw data has no geospatial fields at all; a small manually-curated coordinate table "
         "(R/city_coords.R) is the only way to satisfy the confirmed India-route-map requirement."),
    ],
)

doc.add_heading("3.4 Skew handling for price", level=2)
doc.add_paragraph(
    "price spans ₹{} to ₹{} (a ~{}× range) and is strongly right-skewed (mean ₹{} vs. median "
    "₹{}, §2.2). High Business-class fares are not data errors — deleting them as “outliers” "
    "would misrepresent the real “how much more does Business cost” question a user might ask. Instead, "
    "chart #1 exposes a live log-scale toggle (scale_y_log10()) so a user can switch between the linear view "
    "(which shows absolute rupee gaps) and the log view (which shows proportional/relative differences and "
    "makes the shape of each class's distribution comparable) — directly applying Lecture 4's guidance on "
    "log-transforming skewed data, as an interactive choice rather than a single fixed decision made on the "
    "user's behalf.".format(fmt(STATS["price_min"]), fmt(STATS["price_max"]),
                             round(STATS["price_max"] / STATS["price_min"]),
                             fmt(STATS["price_mean"]), fmt(STATS["price_median"]))
)

doc.add_page_break()

# ============================================================================
# 4. VISUALIZATION & INTERACTION DESIGN
# ============================================================================
doc.add_heading("4. Visualization & Interaction Design", level=1)

doc.add_paragraph(
    "Nine charts were built, deliberately spanning the task-based chart directory taught in Lecture 9 "
    "(Amounts, Distributions, Proportions, x-y relationships, and a geospatial view), and every encoding choice "
    "below is justified against a specific piece of course theory rather than aesthetic preference alone."
)

doc.add_heading("4.1 Chart-by-chart justification", level=2)

charts = [
    ("Chart #1 — Price distribution by class", "fig1_price_by_class.png",
     "Violin + jittered points, Traveler tab.",
     "A violin plot was chosen over a boxplot because the Economy/Business price distribution is genuinely "
     "bimodal (§2.3) — a boxplot's five-number summary would compress two distinct populations into a "
     "single box and hide the bimodality entirely. Lecture 9 explicitly recommends violin/sina/ridgeline plots "
     "over boxplots for exactly this case. The log-scale toggle (§3.4) is wired to this chart."),
    ("Chart #2 — Average price by airline", "fig2_avg_price_airline.png",
     "Ordered horizontal bar, Airline Analytics tab.",
     "Bar length encodes position on a common scale — Cleveland & McGill's empirically most accurate "
     "graphical-perception encoding (Lecture 2), and Stevens' Power Law gives bar length a perceptual exponent "
     "of ~1.0 (i.e. perceived magnitude matches actual magnitude almost exactly; Lecture 3). Bars are sorted by "
     "value, not alphabetically, for immediate rank readability."),
    ("Chart #3 — Market share by airline", "fig3_market_share.png",
     "Ordered bar (deliberately NOT a pie chart), Airline Analytics tab.",
     "See §4.3 for the full rejection rationale. In short: 6 categories with no extreme skew is exactly the "
     "regime where Cleveland & McGill's ranking (position/length > angle/area) predicts real ranking errors from "
     "a pie chart; a bar chart communicates the same “share” information with a provably more accurate "
     "encoding."),
    ("Chart #4 — Price vs. days-left booking curve", "fig4_booking_curve.png",
     "Mean price per days_left, per class, with a LOESS trend line, Traveler tab.",
     "Directly answers “when should I book?” (§2.4). Data is aggregated to (days_left, class) means "
     "before smoothing — fitting LOESS directly on the ~300k raw rows is computationally impractical (LOESS "
     "scales very poorly with n) and days_left only has 49 distinct values, so aggregating first is both faster "
     "and, arguably, a cleaner signal than a smooth fit over a raw point cloud. This follows Lecture 4's taught "
     "LOESS/geom_smooth technique for revealing trend in noisy data."),
    ("Chart #5 — Route heatmap", "fig5_route_heatmap.png",
     "source_city × destination_city grid, fill = mean price or flight count (Reconfigure toggle), "
     "Airline Analytics tab.",
     "Lecture 9 notes heatmaps are good for spotting broad trend but weak for precise value readout; in the "
     "live app this is compensated for with a plotly hover tooltip giving the exact number on demand "
     "(Shneiderman's “details-on-demand”, Lecture 2). The static image above therefore also prints the "
     "value directly on each cell, since a printed report page has no hover interaction."),
    ("Chart #6 — India route map", "fig6_india_map.png",
     "Geospatial flow map: 6 cities as points, routes as weighted line segments, Airline Analytics tab.",
     "The dataset covers only 6 Indian cities (domestic routes) — a world map would misrepresent the data's "
     "actual geographic scope, so the map is deliberately India-only. In the live app this is an interactive "
     "leaflet map (line width/color encode flight volume or price, a flat 2D length/color encoding rather than "
     "a distorting 3D or bubble-volume encoding); the image above is a static ggplot equivalent generated purely "
     "for this report, since a printed page cannot embed the interactive leaflet widget."),
    ("Chart #7 — Duration vs. price scatter", "fig7_duration_price_scatter.png",
     "Scatterplot colored by stops (or class / airline), Traveler tab, brushable in the live app.",
     "A scatter (not a bubble chart) was used deliberately: Lecture 9 explicitly criticizes bubble charts for "
     "confounding two variables (position and size) in one mark, hurting perceptual accuracy. Color-by-stops "
     "uses a qualitative Bertin visual variable appropriately, since stops has only 3 ordered categories. In the "
     "live app, points are capped to a 5,000-point random sample for interactive rendering performance (the "
     "unfiltered dataset has ~300k rows, far more than a browser can smoothly render as individually "
     "hoverable/selectable marks); the sample preserves the overall shape of the relationship while keeping the "
     "chart responsive."),
    ("Chart #8 — Departure-time × airline grid", "fig8_time_airline_grid.png",
     "Small-multiple heatmap, rows ordered by time-of-day (not alphabetically, §3.2), Traveler tab.",
     "Demonstrates multivariate encoding via small multiples (Bertin) and is the chart most directly used to "
     "showcase the Connect/linked-brushing interaction (§4.4): flights lasso-selected on chart #7 are "
     "highlighted here as an overlay, letting a user see where their selection concentrates across time-of-day "
     "and airline."),
]

for title_text, fname, meta_text, justification in charts:
    doc.add_heading(title_text, level=3)
    p = doc.add_paragraph()
    p.add_run(meta_text).italic = True
    add_figure(fname, f"{title_text} — static export for this report.")
    doc.add_paragraph(justification)

doc.add_heading("Chart #9 — KPI tiles", level=3)
doc.add_paragraph(
    "Four stat tiles (flight count, average price, average duration, average price-per-hour on the Traveler "
    "tab; flight count, market leader, cheapest-average airline, average price on the Airline tab) sit above "
    "the detail charts on both tabs, always reflecting the current filter state. This is Stephen Few's "
    "single-screen dashboard principle (Lecture 8) applied directly: an at-a-glance overview before any "
    "zooming/filtering, per Shneiderman's mantra."
)
placeholder("Screenshot of the live KPI tiles row from the running app (both tabs)")

doc.add_heading("4.2 The pie-chart decision", level=2)
doc.add_paragraph(
    "No pie chart is used anywhere in the app, despite airline (6 categories) and class/stops (2–3 "
    "categories) being exactly the kind of “simple category share” situation where a pie chart is "
    "tempting. Cleveland & McGill's perceptual-accuracy ranking (position on a common scale > position "
    "non-aligned > length/angle > area > volume > shading, Lecture 2) and Stevens' Power Law (Lecture 3: bar "
    "length exponent ≈1.0/accurate vs. area exponent ≈0.7/underestimated) both rank angle/area judgments "
    "— exactly what a pie chart requires — below simple length judgments. With 6 non-extreme-share "
    "airline categories (§2.3's market-share table shows shares ranging 3%–43%, several fairly close "
    "together), this is precisely the regime where pie-slice angle misjudgment causes real ranking errors. "
    "Chart #3 uses an ordered bar instead, communicating the same information with a provably more accurate "
    "encoding. This decision was made explicitly, not by default — see §5 for how an AI-suggested pie "
    "chart was reviewed and rejected during development."
)

doc.add_heading("4.3 Interaction design (9 operators)", level=2)
add_table(
    ["Operator", "Concrete implementation"],
    [
        ("Filter", "Sidebar sliderInput (days_left, duration), checkboxGroupInput/selectInput (airline, class, "
                    "stops, source/dest city) feeding one shared reactive() filtered dataset per tab."),
        ("Select", "plotly event_data(\"plotly_click\"/\"plotly_selected\") on the route heatmap and the scatter."),
        ("Explore / Change configuration (continuous)", "Slider drag → live re-render of every chart on the tab."),
        ("Reconfigure", "radioButtons metric toggle (Mean Price / Flight Count) re-fills the heatmap and the "
                         "time×airline grid by a different measure without changing the chart type."),
        ("Encode", "Log-scale checkbox on chart #1; color-by selector (stops/class/airline) on chart #7."),
        ("Abstract/Elaborate", "KPI tiles (abstract, tab-wide) sit above the detail charts — overview-first, "
                                "detail-on-demand."),
        ("Connect (linked brushing)", "A lasso selection on chart #7 highlights the matching flights as an "
                                       "overlay on chart #8; a click on chart #5 highlights that route on chart #6."),
        ("Undo/Redo", "“Reset all filters” button per tab restores every input to its default (a scoped, "
                       "single-step interpretation of the operator, noted here explicitly rather than a full "
                       "history stack)."),
        ("Change configuration (persona)", "The navbarPage tab switch between Traveler View and Airline "
                                            "Analytics changes the entire chart configuration/persona context."),
    ],
)

doc.add_heading("4.4 Dashboard layout", level=2)
doc.add_paragraph(
    "Both tabs use a fixed sidebar (~25% width, filters + reset button) and a main panel (~75%, KPI row on top, "
    "then two rows of two charts) with no vertical scroll-chasing and no decorative gradients or borders — "
    "Stephen Few's single-screen, high-data-ink-ratio dashboard principle (Lecture 8). A single colorblind-safe "
    "qualitative palette (RColorBrewer “Set2”) is used consistently for airline across both tabs, so "
    "the same airline is always the same color everywhere it appears (Lecture 11: consistent, accessible color "
    "mapping)."
)
placeholder("Full-page screenshot of the Traveler View tab (sidebar + KPI tiles + 4 charts)")
placeholder("Full-page screenshot of the Airline Analytics tab (sidebar + KPI tiles + 4 charts)")

doc.add_page_break()

# ============================================================================
# 5. AI USAGE DOCUMENTATION
# ============================================================================
doc.add_heading("5. AI Usage Documentation", level=1)
doc.add_paragraph(
    "The full, running prompt-by-prompt log is kept in report/ai_prompt_appendix.md and is updated by each "
    "team member as they work; this section summarizes it. Claude Code (Anthropic) was the AI tool used "
    "throughout the initial planning and implementation phase covered below; teammates should add their own "
    "tool(s) (ChatGPT, Copilot, etc.) and entries to the appendix as they contribute."
)

doc.add_heading("5.1 Capabilities used", level=2)
add_table(
    ["Capability", "Where it was used"],
    [
        ("Idea generation / research", "Surveying all 11 lecture PDFs to identify which taught theory (Bertin, "
                                        "Cleveland-McGill, Stevens' Power Law, Tufte, the 9 interaction operators, "
                                        "Stephen Few, Pyramid Principle) should ground each design decision."),
        ("Code generation", "The full R Shiny app — data_prep.R, all 9 chart functions, ui.R, server.R, "
                             "app.R — plus the report-asset generation scripts."),
        ("Debugging", "Diagnosing and fixing 3 concrete bugs found during verification (§6.2)."),
        ("Data analysis", "Running the EDA and uncovering the class-mix confound described in §2.3."),
        ("Visualization suggestions", "Proposing and justifying each of the 9 chart types against course theory."),
        ("UI design", "The two-tab, sidebar-plus-KPI-tiles dashboard layout (§4.4)."),
        ("Refactoring", "Extracting the shared theme/palette into R/theme.R so all 9 charts look like one system."),
        ("Documentation", "PLAN.md, README.md, and this report."),
    ],
)

doc.add_heading("5.2 How AI outputs were validated or corrected", level=2)
doc.add_paragraph(
    "AI-generated code and suggestions were not accepted uncritically. Concrete examples from this project:"
)
validation_items = [
    "A first instinct to visualize market share with a pie chart was reconsidered and rejected in favor of an "
    "ordered bar chart, specifically because Cleveland-McGill/Stevens' theory (Lecture 2/3) predicts real "
    "misjudgment for 6 non-extreme-share categories — see §4.2.",
    "The scatter plot's linked-brushing selection was initially implemented using plotly's raw pointNumber as a "
    "row index; this was caught as incorrect once tested, because plotly resets pointNumber per color-trace when "
    "a categorical color mapping splits the chart into multiple traces. It was fixed by adding a stable row-id "
    "“key” aesthetic instead — caught by writing and running a headless smoke test, not assumed "
    "correct from the code alone.",
    "The heatmap-click and scatter-lasso interactions initially had no effect at all when tested with "
    "shiny::testServer() — plotly's event_data() silently returns nothing unless the chart explicitly calls "
    "event_register() for that event type. This is not obvious from reading the code and was only caught by "
    "actually exercising the reactive server logic programmatically.",
    "The booking-curve chart originally fit a LOESS smoother directly on the full ~300,000-row dataset; a live "
    "test run of the app hung during rendering. This was diagnosed as a genuine performance/scalability bug "
    "(LOESS scales very poorly with n) and fixed by aggregating to per-days_left means before smoothing, rather "
    "than papering over the symptom.",
    "The duration-range slider's default bounds used round(), which rounded the true maximum duration down and "
    "silently excluded a small number of legitimate boundary flights even at the slider's default "
    "“full range” position. This was only caught because a headless test explicitly asserted that the "
    "default filter state should return every row, and it didn't — fixed with floor()/ceiling() instead.",
]
for item in validation_items:
    doc.add_paragraph(item, style="List Bullet")

placeholder("Add each team member's own AI usage entries here (tool, prompt, capability, validation) as they "
            "complete their assigned section — see report/ai_prompt_appendix.md for the template")

doc.add_page_break()

# ============================================================================
# 6. EVALUATION & REFLECTION
# ============================================================================
doc.add_heading("6. Evaluation & Reflection", level=1)

doc.add_heading("6.1 Strengths", level=2)
strengths = [
    "Dual-persona coverage: the same underlying data serves two genuinely different decision-support goals "
    "(traveler vs. airline analyst) without duplicating the data pipeline.",
    "Every encoding choice is justified against specific, citable course theory rather than default chart-library "
    "behavior — including a deliberate, reasoned rejection of a pie chart.",
    "The class-mix confound found in §2.3 is a genuine, non-obvious insight (not a canned dataset fact) that "
    "directly shaped why the class filter is a first-class, always-visible control rather than a secondary option.",
    "The tech stack (R + Shiny + ggplot2/plotly/leaflet) mirrors the course's own worked Lecture 6 example almost "
    "exactly, making the implementation easy to map back to what was taught.",
    "Headless verification (a chart-function smoke test and a full shiny::testServer() run of the reactive graph) "
    "caught real bugs before manual testing, rather than relying on visual inspection alone.",
]
for s in strengths:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("6.2 Limitations", level=2)
limitations = [
    "No real calendar dates — only days_left (booking lead time) — so no seasonal/day-of-week pricing "
    "patterns can be shown, only lead-time patterns.",
    "No flight-delay, cancellation, or on-time-performance data, so the “best airline” comparison is "
    "price/duration-only, not reliability-aware.",
    "The India route map is schematic (straight-line arcs between manually-curated city centroids), not a true "
    "flight-path or air-corridor visualization.",
    "This is a static snapshot of historical listings, not live/real-time pricing — the dashboard supports "
    "retrospective exploration, not real-time booking decisions.",
    "No predictive element by design (see §1.3) — the dashboard cannot forecast a specific future price, "
    "only show historical patterns by booking lead time.",
    "Formal usability evaluation (think-aloud testing, expert heuristic review — Lecture 11) was not "
    "performed due to project time constraints; verification here was technical (does the code work correctly) "
    "rather than user-facing (do real users find the intended insights).",
]
for l in limitations:
    doc.add_paragraph(l, style="List Bullet")

doc.add_heading("6.3 What AI helped with vs. human decisions", level=2)
doc.add_paragraph(
    "AI accelerated: boilerplate Shiny wiring, first-draft chart code, the initial course-theory research pass "
    "across all 11 lectures, and this report's drafting. Human/team decisions (to be confirmed and adjusted by "
    "the team during review): which two personas to build around, the India-only map scope, the decision to "
    "exclude a predictive model, and — per §5.2 — catching and rejecting a suggested pie chart in "
    "favor of theory-backed alternatives."
)
placeholder("Team: replace/expand this subsection with your own account of which decisions were genuinely "
            "human judgment calls versus accepted AI suggestions, once you've reviewed the app together")

doc.add_heading("6.4 Lessons learned and future improvements", level=2)
lessons = [
    "Never trust an unconditioned aggregate statistic without checking for a lurking categorical variable "
    "(§2.3's class-mix confound is this project's own Anscombe's Quartet moment) — this is now baked "
    "into the app itself via the always-visible class filter, not just a report footnote.",
    "If a real, dated booking-price time series became available, chart #4 could show actual seasonal trends "
    "instead of only lead-time trends.",
    "Linked brushing (§4.3's Connect operator) is currently implemented between one chart pair per tab; a "
    "future iteration could extend it across all charts on a tab simultaneously.",
    "A structured usability pass (Lecture 11: think-aloud test with 2–3 people outside the team) would "
    "validate whether the intended insights (e.g. the booking-curve trend, the class-mix finding) are actually "
    "discoverable by a first-time user without guidance.",
]
for l in lessons:
    doc.add_paragraph(l, style="List Bullet")

doc.add_page_break()

# ============================================================================
# 7. DELIVERABLES CHECKLIST
# ============================================================================
doc.add_heading("7. Deliverables Checklist", level=1)
add_table(
    ["Deliverable", "Status", "Location"],
    [
        ("Report", "Complete (this document)", "Report.docx"),
        ("Source code", "Complete and verified", "GitHub: Eyalmied/Visualization-Project"),
        ("AI prompt appendix", "Template + first entries in place; add your own as you work",
         "report/ai_prompt_appendix.md"),
        ("Working demo (10–15 min)", "App built and verified; live walkthrough not yet rehearsed as a team",
         "shiny::runApp() from the repo root"),
    ],
)
doc.add_paragraph(
    "Suggested demo structure (Pyramid Principle, Lecture 11): open with the headline finding — e.g. "
    "“booking 30+ days out saves about {}% versus booking in the final 3 days, and that gap looks even "
    "bigger or smaller depending on which airline and class you pick” — then let each team member "
    "walk through their assigned tab/charts as supporting evidence, finishing with a live interactive "
    "demonstration of the filters and linked-brushing.".format(STATS["booking_pct_diff"])
)

out_path = ROOT / "Report.docx"
doc.save(str(out_path))
print(f"Saved {out_path}")
